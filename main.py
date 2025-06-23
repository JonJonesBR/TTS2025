from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import os
import io
import asyncio
import edge_tts
from PyPDF2 import PdfReader
from docx import Document
from ebooklib import epub
from bs4 import BeautifulSoup
import traceback
import json
import uuid
import re
import unicodedata
import subprocess
from pathlib import Path

import nest_asyncio
import aiohttp

from num2words import num2words
import chardet
import html2text

nest_asyncio.apply()

app = FastAPI()

# NOVO: Monta StaticFiles para servir arquivos do diretório 'static' na raiz
# Isso fará com que '/index.html' seja servido por padrão em '/'
app.mount("/", StaticFiles(directory="static", html=True), name="static_root")

# As rotas /voices, /process_file, /status, /download e /set_gemini_api_key
# continuarão a funcionar normalmente.

cached_voices = {}
conversion_tasks = {}
GEMINI_API_KEY = None

FFMPEG_BIN = "ffmpeg"

# ================== CONFIGURAÇÕES E MAPAS PARA LIMPEZA DE TEXTO ==================

ABREVIACOES_MAP = {
    'dr': 'Doutor', 'd': 'Dona', 'dra': 'Doutora',
    'sr': 'Senhor', 'sra': 'Senhora', 'srta': 'Senhorita',
    'prof': 'Professor', 'profa': 'Professora',
    'eng': 'Engenheiro', 'engª': 'Engenheira',
    'adm': 'Administrador', 'adv': 'Advogado',
    'exmo': 'Excelentíssimo', 'exma': 'Excelentíssima',
    'v.exa': 'Vossa Excelência', 'v.sa': 'Vossa Senhoria',
    'av': 'Avenida', 'r': 'Rua', 'km': 'Quilômetro',
    'etc': 'etcétera', 'ref': 'Referência',
    'pag': 'Página', 'pags': 'Páginas',
    'fl': 'Folha', 'fls': 'Folhas',
    'pe': 'Padre',
    'dept': 'Departamento', 'depto': 'Departamento',
    'univ': 'Universidade', 'inst': 'Instituição',
    'est': 'Estado', 'tel': 'Telefone',
    'eua': 'Estados Unidos da América',
    'ed': 'Edição', 'ltda': 'Limitada'
}

ABREVIACOES_MAP_LOWER = {k.lower(): v for k, v in ABREVIACOES_MAP.items()}

CASOS_ESPECIAIS_RE = {
     r'\bV\.Exa\.(?=\s)': 'Vossa Excelência',
     r'\bV\.Sa\.(?=\s)': 'Vossa Senhoria',
     r'\bEngª\.(?=\s)': 'Engenheira'
}

CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM = {
    'UM': '1', 'DOIS': '2', 'TRÊS': '3', 'QUATRO': '4', 'CINCO': '5',
    'SEIS': '6', 'SETE': '7', 'OITO': '8', 'NOVE': '9', 'DEZ': '10',
    'ONZE': '11', 'DOZE': '12', 'TREZE': '13', 'CATORZE': '14', 'QUINZE': '15',
    'DEZESSEIS': '16', 'DEZESSETE': '17', 'DEZOITO': '18', 'DEZENOVE': '19', 'VINTE': '20'
}

ABREVIACOES_QUE_NAO_TERMINAM_FRASE = set([
    'sr.', 'sra.', 'srta.', 'dr.', 'dra.', 'prof.', 'profa.', 'eng.', 'exmo.', 'exma.',
    'pe.', 'rev.', 'ilmo.', 'ilma.', 'gen.', 'cel.', 'maj.', 'cap.', 'ten.', 'sgt.',
    'cb.', 'sd.', 'me.', 'ms.', 'msc.', 'esp.', 'av.', 'r.', 'pç.', 'esq.', 'trav.',
    'jd.', 'pq.', 'rod.', 'km.', 'apt.', 'ap.', 'bl.', 'cj.', 'cs.', 'ed.', 'nº',
    'no.', 'uf.', 'cep.', 'est.', 'mun.', 'dist.', 'zon.', 'reg.', 'kg.', 'cm.',
    'mm.', 'lt.', 'ml.', 'mg.', 'seg.', 'min.', 'hr.', 'ltda.', 's.a.', 's/a',
    'cnpj.', 'cpf.', 'rg.', 'proc.', 'ref.', 'cod.', 'tel.', 'etc.', 'p.ex.', 'ex.',
    'i.e.', 'e.g.', 'vs.', 'cf.', 'op.cit.', 'loc.cit.', 'fl.', 'fls.', 'pag.',
    'p.', 'pp.', 'u.s.', 'e.u.a.', 'o.n.u.', 'i.b.m.', 'h.p.', 'obs.', 'att.',
    'resp.', 'publ.', 'ed.', 'doutora', 'senhora', 'senhor', 'doutor', 'professor',
    'professora', 'general'
])
SIGLA_COM_PONTOS_RE = re.compile(r'\b([A-Z]\.\s*)+$')

# ================== FUNÇÕES AUXILIARES DE LIMPEZA DE TEXTO ==================

def _formatar_numeracao_capitulos(texto):
    def substituir_cap(match):
        tipo_cap = match.group(1).upper()
        numero_rom_arab = match.group(2)
        numero_extenso = match.group(3)
        titulo_opcional = match.group(4).strip() if match.group(4) else ""

        numero_final = ""
        if numero_rom_arab:
            numero_final = numero_rom_arab.upper()
        elif numero_extenso:
            num_ext_upper = numero_extenso.strip().upper()
            numero_final = CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM.get(num_ext_upper, num_ext_upper)

        cabecalho = f"{tipo_cap} {numero_final}."
        if titulo_opcional:
            palavras_titulo = []
            for p in titulo_opcional.split():
                if p.isupper() and len(p) > 1:
                    palavras_titulo.append(p)
                else:
                    palavras_titulo.append(p.capitalize())
            titulo_formatado = " ".join(palavras_titulo)
            return f"\n\n{cabecalho}\n\n{titulo_formatado}"
        return f"\n\n{cabecalho}\n\n"

    padrao = re.compile(
        r'(?i)(cap[íi]tulo|cap\.?)\s+'
        r'(?:(\d+|[IVXLCDM]+)|([A-ZÇÉÊÓÃÕa-zçéêóãõ]+))'
        r'\s*[:\-.]?\s*'
        r'(?=\S)([^\n]*)?',
        re.IGNORECASE
    )
    texto = padrao.sub(substituir_cap, texto)

    def substituir_extenso_com_titulo(match):
        num_ext = match.group(1).strip().upper()
        titulo = match.group(2).strip().title()
        numero = CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM.get(num_ext, num_ext)
        return f"CAPÍTULO {numero}: {titulo}"

    padrao_extenso_titulo = re.compile(r'CAP[IÍ]TULO\s+([A-ZÇÉÊÓÃÕ]+)\s*[:\-]\s*(.+)', re.IGNORECASE)
    texto = padrao_extenso_titulo.sub(substituir_extenso_com_titulo, texto)
    return texto

def _remover_numeros_pagina_isolados(texto):
    linhas = texto.splitlines()
    novas_linhas = []
    for linha in linhas:
        if re.match(r'^\s*\d+\s*$', linha):
            continue
        linha = re.sub(r'\s{3,}\d+\s*$', '', linha)
        novas_linhas.append(linha)
    return '\n'.join(novas_linhas)

def _normalizar_caixa_alta_linhas(texto):
    linhas = texto.splitlines()
    texto_final = []
    for linha in linhas:
        if not re.match(r'^\s*CAP[ÍI]TULO\s+[\w\d]+\.?\s*$', linha, re.IGNORECASE):
            if linha.isupper() and len(linha.strip()) > 3 and any(c.isalpha() for c in linha):
                palavras = []
                for p in linha.split():
                    if len(p) > 1 and p.isupper() and p.isalpha() and p not in ['I', 'A', 'E', 'O', 'U']:
                        if not (sum(1 for char in p if char in "AEIOU") > 0 and \
                                sum(1 for char in p if char not in "AEIOU") > 0 and len(p) <=4) :
                            palavras.append(p)
                            continue
                    palavras.append(p.capitalize())
                texto_final.append(" ".join(palavras))
            else:
                texto_final.append(linha)
        else:
            texto_final.append(linha)
    return "\n".join(texto_final)

def _corrigir_hifenizacao_quebras(texto):
    return re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', texto)

def _remover_metadados_pdf(texto):
    texto = re.sub(r'^\s*[\w\d_-]+\.indd\s+\d+\s+\d{2}/\d{2}/\d{2,4}\s+\d{1,2}:\d{2}(:\d{2})?\s*([AP]M)?\s*$', '', texto, flags=re.MULTILINE)
    return texto

def _expandir_abreviacoes_numeros(texto: str) -> str:
    """Expande abreviações comuns (removendo o ponto da abrev.) e converte números."""

    for abrev_re, expansao in CASOS_ESPECIAIS_RE.items():
         texto = re.sub(abrev_re, expansao, texto, flags=re.IGNORECASE)

    def replace_abrev_com_ponto(match):
        abrev_encontrada = match.group(1)
        expansao = ABREVIACOES_MAP_LOWER.get(abrev_encontrada.lower())
        if expansao:
            return expansao
        else:
            return match.group(0)

    chaves_escapadas = [re.escape(k) for k in ABREVIACOES_MAP_LOWER.keys() if '.' not in k and 'ª' not in k]
    if chaves_escapadas:
        padrao_abrev_simples = r'\b(' + '|'.join(chaves_escapadas) + r')\.'
        texto = re.sub(padrao_abrev_simples, replace_abrev_com_ponto, texto, flags=re.IGNORECASE)

    # Conversão de números cardinais
    def _converter_numero_match(match):
        num_str = match.group(0)
        try:
            if re.match(r'^\d{4}$', num_str) and (1900 <= int(num_str) <= 2100): return num_str
            if len(num_str) > 7 : return num_str
            return num2words(int(num_str), lang='pt_BR')
        except Exception: return num_str
    texto = re.sub(r'\b\d+\b', _converter_numero_match, texto)

    # Conversão de valores monetários
    def _converter_valor_monetario_match(match):
        valor_inteiro = match.group(1).replace('.', '')
        try: return f"{num2words(int(valor_inteiro), lang='pt_BR')} reais"
        except Exception: return match.group(0)
    texto = re.sub(r'R\$\s*(\d{1,3}(?:\.\d{3})*),(\d{2})', _converter_valor_monetario_match, texto)
    texto = re.sub(r'R\$\s*(\d+)(?:,00)?', lambda m: f"{num2words(int(m.group(1)), lang='pt_BR')} reais" if m.group(1) else m.group(0) , texto)

    # Conversão de intervalos numéricos
    texto = re.sub(r'\b(\d+)\s*-\s*(\d+)\b', lambda m: f"{num2words(int(m.group(1)), lang='pt_BR')} a {num2words(int(m.group(2)), lang='pt_BR')}", texto)
    
    return texto

def _converter_ordinais_para_extenso(texto: str) -> str:
    """Converte números ordinais como 1º, 2a, 3ª para extenso."""

    def substituir_ordinal(match):
        numero = match.group(1)
        terminacao = match.group(2).lower()

        try:
            num_int = int(numero)
            if terminacao == 'o' or terminacao == 'º':
                return num2words(num_int, lang='pt_BR', to='ordinal')
            elif terminacao == 'a' or terminacao == 'ª':
                ordinal_masc = num2words(num_int, lang='pt_BR', to='ordinal')
                if ordinal_masc.endswith('o'):
                    return ordinal_masc[:-1] + 'a'
                else:
                    return ordinal_masc
            else:
                return match.group(0)
        except ValueError:
            return match.group(0)

    padrao_ordinal = re.compile(r'\b(\d+)\s*([oaºª])(?!\w)', re.IGNORECASE)
    texto = padrao_ordinal.sub(substituir_ordinal, texto)

    return texto

# ================== FUNÇÃO PRINCIPAL DE FORMATAÇÃO ==================
def formatar_texto_para_tts(texto_bruto: str) -> str:
    print("⚙️ Aplicando formatações ao texto para TTS...")
    texto = texto_bruto

    # 0. Normalizações e remoções básicas
    texto = unicodedata.normalize('NFKC', texto)
    texto = texto.replace('\f', '\n\n')
    texto = texto.replace('*', '')
    caracteres_para_espaco = ['_', '#', '@']
    caracteres_para_remover = ['(', ')', '\\', '[', ']']
    for char in caracteres_para_espaco: texto = texto.replace(char, ' ')
    for char in caracteres_para_remover: texto = texto.replace(char, '')
    texto = re.sub(r'\{.*?\}', '', texto)

    # 1. Pré-limpeza de espaços múltiplos e linhas vazias
    texto = re.sub(r'[ \t]+', ' ', texto)
    texto = "\n".join([linha.strip() for linha in texto.splitlines() if linha.strip()])

    # 2. JUNTAR LINHAS DENTRO DE PARÁGRAFOS INTENCIONAIS
    paragrafos_originais = texto.split('\n\n')
    paragrafos_processados = []
    for paragrafo_bruto in paragrafos_originais:
        paragrafo_bruto = paragrafo_bruto.strip()
        if not paragrafo_bruto:
            continue
        linhas_do_paragrafo = paragrafo_bruto.split('\n')
        buffer_linha_atual = ""
        for i, linha in enumerate(linhas_do_paragrafo):
            linha_strip = linha.strip()
            if not linha_strip:
                continue
            juntar_com_anterior = False
            if buffer_linha_atual:
                ultima_palavra_buffer = buffer_linha_atual.split()[-1].lower() if buffer_linha_atual else ""
                termina_abreviacao = ultima_palavra_buffer in ABREVIACOES_QUE_NAO_TERMINAM_FRASE
                termina_sigla_ponto = re.search(r'\b[A-Z]\.$', buffer_linha_atual) is not None
                termina_pontuacao_forte = re.search(r'[.!?…]$', buffer_linha_atual)
                nao_juntar = False
                if termina_pontuacao_forte and not termina_abreviacao and not termina_sigla_ponto:
                     if linha_strip and linha_strip[0].isupper(): nao_juntar = True
                if termina_abreviacao or termina_sigla_ponto: juntar_com_anterior = True
                elif not nao_juntar and not termina_pontuacao_forte: juntar_com_anterior = True
                elif buffer_linha_atual.lower() in ['doutora', 'senhora', 'senhor', 'doutor']: juntar_com_anterior = True
            if juntar_com_anterior: buffer_linha_atual += " " + linha_strip
            else:
                if buffer_linha_atual: paragrafos_processados.append(buffer_linha_atual)
                buffer_linha_atual = linha_strip
        if buffer_linha_atual: paragrafos_processados.append(buffer_linha_atual)
    texto = '\n\n'.join(paragrafos_processados)
    
    # 3. Limpeza de espaços e quebras
    texto = re.sub(r'[ \t]+', ' ', texto)
    texto = re.sub(r'(?<!\n)\n(?!\n)', ' ', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)

    # 4. Formatações que operam melhor no texto mais estruturado
    texto = _remover_metadados_pdf(texto)
    texto = _remover_numeros_pagina_isolados(texto)
    texto = _corrigir_hifenizacao_quebras(texto)
    texto = _formatar_numeracao_capitulos(texto)

    # 5. REINTRODUZIR QUEBRAS DE PARÁGRAFO (\n\n) INTELIGENTEMENTE
    segmentos = re.split(r'([.!?…])\s*', texto)
    texto_reconstruido = ""; buffer_segmento = ""
    for i in range(0, len(segmentos), 2):
        parte_texto = segmentos[i]
        pontuacao = segmentos[i+1] if i + 1 < len(segmentos) else ""
        segmento_completo = (parte_texto + pontuacao).strip()
        if not segmento_completo: continue
        ultima_palavra = segmento_completo.split()[-1].lower() if segmento_completo else ""
        ultima_palavra_sem_ponto = ultima_palavra.rstrip('.!?…') if pontuacao else ultima_palavra
        termina_abreviacao_conhecida = ultima_palavra in ABREVIACOES_QUE_NAO_TERMINAM_FRASE or \
                                        ultima_palavra_sem_ponto in ABREVIACOES_QUE_NAO_TERMINAM_FRASE
        termina_sigla_padrao = SIGLA_COM_PONTOS_RE.search(segmento_completo) is not None
        nao_quebrar = False
        if pontuacao == '.':
             if termina_abreviacao_conhecida or termina_sigla_padrao: nao_quebrar = True
        if buffer_segmento: buffer_segmento += " " + segmento_completo
        else: buffer_segmento = segmento_completo
        if not nao_quebrar: texto_reconstruido += buffer_segmento + "\n\n"; buffer_segmento = ""
    if buffer_segmento:
         texto_reconstruido += buffer_segmento
         if not re.search(r'[.!?…)]$', buffer_segmento): texto_reconstruido += "."
         texto_reconstruido += "\n\n"
    texto = texto_reconstruido.strip()

    # 6. Formatações Finais (Caixa, Ordinais, Cardinais, etc.)
    texto = _normalizar_caixa_alta_linhas(texto)
    texto = _converter_ordinais_para_extenso(texto)
    texto = _expandir_abreviacoes_numeros(texto)

    # NOVA ETAPA 6.5: Limpeza Pós-Expansão
    formas_expandidas_tratamento = ['Senhor', 'Senhora', 'Doutor', 'Doutora', 'Professor', 'Professora', 'Excelentíssimo', 'Excelentíssima']
    for forma in formas_expandidas_tratamento:
        padrao_limpeza = r'\b' + re.escape(forma) + r'\.\s+([A-Z])'
        texto = re.sub(padrao_limpeza, rf'{forma} \1', texto)
        padrao_limpeza_sem_espaco = r'\b' + re.escape(forma) + r'\.([A-Z])'
        texto = re.sub(padrao_limpeza_sem_espaco, rf'{forma} \1', texto)
        
    texto = re.sub(r'\b([A-Z])\.\s+([A-Z])', r'\1. \2', texto)
    texto = re.sub(r'\b([A-Z])\.\s+([A-Z][a-z])', r'\1. \2', texto)


    # 7. Limpeza Final de Parágrafos Vazios e Espaços
    paragrafos_finais = texto.split('\n\n')
    paragrafos_formatados_final = []
    for p in paragrafos_finais:
        p_strip = p.strip()
        if not p_strip: continue
        if not re.search(r'[.!?…)]$', p_strip) and \
           not re.match(r'^\s*CAP[ÍI]TULO\s+[\w\d]+\.?\s*$', p_strip.split('\n')[0].strip(), re.IGNORECASE):
            p_strip += '.'
        paragrafos_formatados_final.append(p_strip)
    texto = '\n\n'.join(paragrafos_formatados_final)
    texto = re.sub(r'[ \t]+', ' ', texto).strip()
    texto = re.sub(r'\n{2,}', '\n\n', texto)

    print("✅ Formatação de texto para TTS concluída.")
    return texto.strip()

# ================== FIM DAS FUNÇÕES E CONSTANTES DE LIMPEZA DE TEXTO ==================


# Função para extrair texto de diferentes tipos de arquivo
async def get_text_from_file(file_path: str, task_id: str):
    text = ""
    filename = os.path.basename(file_path)
    total_parts = 1

    try:
        if filename.endswith('.pdf'):
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                extracted_page_text = page.extract_text()
                if extracted_page_text:
                    text += extracted_page_text + "\n"
                progress = int(((i + 1) / total_pages) * 50)
                conversion_tasks[task_id].update({"progress": progress, "message": f"Extraindo texto de PDF (Página {i+1}/{total_pages})..."})
                await asyncio.sleep(0.01)

        elif filename.endswith('.txt'):
            raw_data = open(file_path, 'rb').read()
            detected_encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            with open(file_path, 'r', encoding=detected_encoding, errors='replace') as f:
                text = f.read()
            conversion_tasks[task_id].update({"progress": 50, "message": "Texto de arquivo TXT lido."})
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                progress = int(((i + 1) / total_paragraphs) * 50)
                conversion_tasks[task_id].update({"progress": progress, "message": f"Extraindo texto de DOCX (Parágrafo {i+1}/{total_paragraphs})..."})
                await asyncio.sleep(0.01)
        elif filename.endswith('.epub'):
            text = _extrair_texto_de_epub_helper(file_path)
            conversion_tasks[task_id].update({"progress": 50, "message": "Texto de arquivo EPUB extraído."})

        conversion_tasks[task_id].update({"progress": 50, "message": "Extração de texto concluída."})
        print(f"Extração de texto para {filename} concluída. Total de caracteres: {len(text)}.")
        return text.strip()
    except Exception as e:
        print(f"Erro na extração de texto de {filename}: {e}")
        conversion_tasks[task_id].update({"status": "failed", "message": f"Erro na extração de texto: {str(e)}"})
        raise

def _extrair_texto_de_epub_helper(caminho_epub: str) -> str:
    texto_completo = ""
    try:
        book = epub.read_epub(caminho_epub)
        document_items = [item for item in book.get_items() if item.get_type() == epub.ITEM_DOCUMENT]

        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.ignore_emphasis = False
        h.body_width = 0

        for item in document_items:
            try:
                html_bytes = item.get_content()
                detected_encoding = chardet.detect(html_bytes)['encoding'] or 'utf-8'
                html_texto = html_bytes.decode(detected_encoding, errors='replace')
                
                soup = BeautifulSoup(html_texto, 'html.parser')
                for tag in soup(['nav', 'header', 'footer', 'style', 'script', 'figure', 'figcaption', 'aside', 'link', 'meta']):
                    tag.decompose()
                
                content_tag = soup.find('body') or soup
                if content_tag:
                    texto_completo += h.handle(str(content_tag)) + "\n\n"
            except Exception as e_file:
                print(f"⚠️ Erro ao processar item EPUB '{item.id}': {e_file}")

        if not texto_completo.strip():
            print("⚠️ Nenhum conteúdo textual extraído do EPUB.")
            return ""
        return texto_completo
    except Exception as e:
        print(f"❌ Erro geral ao processar EPUB '{caminho_epub}': {e}")
        return ""


async def get_available_voices():
    global cached_voices
    if cached_voices:
        return cached_voices

    print("Buscando vozes Edge TTS disponíveis...")
    try:
        voices = await edge_tts.list_voices()
        pt_br_voices = {}
        for voice in voices:
            if voice["Locale"] == "pt-BR":
                name = voice["ShortName"].replace("pt-BR-", "")
                name = name.replace("Neural", " (Neural)")
                if voice["Gender"] == "Female":
                    name = f"{name} (Feminina)"
                elif voice["Gender"] == "Male":
                    name = f"{name} (Masculina)"
                pt_br_voices[voice["ShortName"]] = name.strip()

        ordered_voices = {}
        if "pt-BR-ThalitaMultilingualNeural" in pt_br_voices:
            ordered_voices["pt-BR-ThalitaMultilingualNeural"] = pt_br_voices["pt-BR-ThalitaMultilingualNeural"]
        if "pt-BR-FranciscaNeural" in pt_br_voices:
            ordered_voices["pt-BR-FranciscaNeural"] = pt_br_voices["pt-BR-FranciscaNeural"]
        if "pt-BR-AntonioNeural" in pt_br_voices:
            ordered_voices["pt-BR-AntonioNeural"] = pt_br_voices["pt-BR-AntonioNeural"]

        for code, name in pt_br_voices.items():
            if code not in ordered_voices:
                ordered_voices[code] = name

        cached_voices = ordered_voices
        print(f"Vozes carregadas: {len(cached_voices)} opções.")
        return cached_voices
    except Exception as e:
        print(f"Erro ao obter vozes Edge TTS: {e}")
        print(traceback.format_exc())
        return {
            "pt-BR-ThalitaMultilingualNeural": "Thalita (Feminina, Neural) - Fallback",
            "pt-BR-FranciscaNeural": "Francisca (Feminina, Neural) - Fallback",
            "pt-BR-AntonioNeural": "Antonio (Masculina, Neural) - Fallback"
        }

# Endpoint de Health Check para o Render
@app.get("/health", response_class=JSONResponse)
async def health_check():
    return {"status": "ok", "message": "Application is healthy."}

# Endpoint para receber a chave API do Gemini
@app.post("/set_gemini_api_key")
async def set_gemini_api_key_endpoint(api_key: str = Form(...)):
    global GEMINI_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="Chave API não pode ser vazia.")
    GEMINI_API_KEY = api_key
    return JSONResponse({"message": "Chave API do Gemini configurada com sucesso!"})

# Função para interagir com a API do Google Gemini para melhoria de texto
async def enhance_text_with_gemini(text: str) -> str:
    prompt = f"""
    Dado o texto de um livro, sua tarefa é revisá-lo e formatá-lo **exclusivamente para a narrativa principal a ser lida por um sistema de Text-to-Speech (TTS) em português do Brasil.**

    **Instruções Essenciais:**
    1.  **Prioridade Total: Focar APENAS na história/conteúdo narrativo principal.**
        * **Remova COMPLETAMENTE:** prefácios, agradecimentos, índices, bibliografias, notas de rodapé extensas, cabeçalhos e rodapés de página, números de página isolados (ex: '1', '2', '3' sem contexto), metadados de PDF, tabelas que não sejam parte da narrativa direta, e quaisquer outras seções que não sejam a história contada ou o conteúdo central do livro.
        * **Mantenha APENAS:** O título do livro (se presente no início do texto narrativo e identificável como tal) e o corpo da história/conteúdo principal. Se o título não for claro ou for parte de uma capa, ignore-o e comece pela narrativa.
    2.  **Correção de Gramática e Ortografia:** Corrija quaisquer erros gramaticais, de concordância e de ortografia.
    3.  **Pontuação Otimizada para Leitura:** Ajuste a pontuação (vírgulas, pontos, etc.) para que o ritmo da leitura TTS seja o mais natural possível, adicionando pausas onde necessário e removendo onde for excessivo.
    4.  **Expansão de Abreviaturas Ambíguas:** Expanda abreviaturas que podem causar confusão na leitura, como "Dr." para "Doutor", "Sra." para "Senhora", "etc." para "etcétera".
    5.  **Normalização de Números e Valores:** Converta números (cardinais e ordinais) para seus equivalentes por extenso (ex: "1" para "um", "2º" para "segundo"). Converta valores monetários (ex: "R$ 10,50" para "dez reais e cinquenta centavos"). Mantenha números grandes (como anos ou números de telefone) em formato numérico se a pronúncia for clara.
    6.  **Remoção de Elementos Visuais:** Remova quaisquer caracteres, símbolos ou formatos que são estritamente visuais e não contribuem para a leitura em áudio (ex: `*`, `_`, `[ ]`, `{{ }}`, links completos, e qualquer outro tipo de formatação que não seja texto legível).
    7.  **Fluxo Contínuo e Natural:** Garanta que o texto flua naturalmente, como se fosse falado por uma pessoa. O texto final não deve conter rupturas abruptas ou elementos que dificultem a fluidez da leitura em voz alta.
    8.  **Manter o Sentido Original:** A edição deve melhorar a legibilidade para TTS, mas não alterar o significado ou a intenção do conteúdo original.
    9.  **Capítulos:** Se houver formatação de capítulos (ex: "CAPÍTULO I", "CAPÍTULO 1", "CAPÍTULO UM"), certifique-se de que a formatação esteja clara e com quebras de parágrafo adequadas (ex: "\n\nCAPÍTULO [NÚMERO/TÍTULO]\n\n").
    10. **Evitar Introduções/Conclusões da IA:** Não adicione nenhuma introdução ("Aqui está o texto melhorado...") ou conclusão ("Espero que isso ajude...") ao seu resultado. Forneça APENAS o texto revisado e formatado.

    Aqui está o texto a ser melhorado:
    ---
    {text}
    ---
    """

    chat_history = []
    chat_history.append({"role": "user", "parts": [{"text": prompt}]})

    payload = {
        "contents": chat_history,
        "generationConfig": {
            "temperature": 0.7,
            "topP": 0.95,
            "topK": 60
        }
    }
    
    try:
        async with aiohttp.ClientSession() as session:
            api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"
            
            async with session.post(api_url, headers={'Content-Type': 'application/json'}, json=payload) as response:
                response.raise_for_status()
                result = await response.json()
                
                if result.get("candidates") and \
                   result["candidates"][0].get("content") and \
                   result["candidates"][0]["content"].get("parts") and \
                   result["candidates"][0]["content"]["parts"][0].get("text"):
                    return result["candidates"][0]["content"]["parts"][0]["text"]
                else:
                    print("API Gemini retornou uma estrutura de resposta inesperada.")
                    print(result)
                    return text
    except aiohttp.ClientResponseError as e:
        print(f"Erro HTTP da API Gemini (Status: {e.status}): {e.message}")
        if e.status == 400: print("Verifique se a chave API é válida ou se o texto contém conteúdo impróprio/grande demais para a IA.")
        return text
    except aiohttp.ClientError as e:
        print(f"Erro de conexão ao chamar a API Gemini: {e}")
        return text
    except Exception as e:
        print(f"Erro inesperado na função enhance_text_with_gemini: {e}")
        print(traceback.format_exc())
        return text

# NOVO: Função para sanitizar o nome do arquivo, removendo caracteres inválidos
def _limpar_nome_arquivo(filename: str) -> str:
    """Remove caracteres inválidos para nomes de arquivo e substitui espaços por underscore."""
    # Remove caracteres inválidos para nomes de arquivo (Windows/Linux/macOS)
    cleaned_name = re.sub(r'[<>:"/\\|?*]', '', filename)
    # Substitui múltiplos espaços e hífens por um único underscore
    cleaned_name = re.sub(r'[\s-]+', '_', cleaned_name)
    # Remove underscores do início ou fim
    cleaned_name = cleaned_name.strip('_')
    # Limita o tamanho do nome do arquivo para evitar problemas de caminho muito longo
    return cleaned_name[:100]

# NOVO: Função para unificar arquivos de áudio temporários usando FFmpeg
def _unificar_audios_ffmpeg(lista_arquivos_temp: list, arquivo_final: str) -> bool:
    """Une arquivos de áudio temporários em um único arquivo final usando FFmpeg."""
    if not lista_arquivos_temp:
        print("⚠️ Nenhum arquivo de áudio para unificar.")
        return False
    
    # Cria um arquivo de lista para o FFmpeg concat demuxer
    # Usar caminhos absolutos para o file list é mais seguro no Render
    dir_saida = os.path.dirname(arquivo_final)
    os.makedirs(dir_saida, exist_ok=True) # Garante que o diretório de saída existe

    # Limpa o nome do arquivo de lista para evitar caracteres problemáticos
    nome_lista_limpo = f"filelist_{uuid.uuid4().hex}.txt" # Nome único para evitar colisões
    lista_txt_path = os.path.join(dir_saida, nome_lista_limpo)

    try:
        with open(lista_txt_path, "w", encoding='utf-8') as f_list:
            for temp_file in lista_arquivos_temp:
                # FFmpeg concat demuxer precisa de caminhos "safe".
                # A melhor prática é usar paths relativos se possível ou escapar.
                # No Render, os arquivos estão no mesmo diretório, então path relativo simples funciona.
                # Ou, para máxima segurança, path absoluto e '-safe 0'. Vamos com absoluto + safe 0.
                safe_path = str(Path(temp_file).resolve()).replace("'", r"\'") # Escapa aspas
                f_list.write(f"file '{safe_path}'\n")
        
        comando = [
            FFMPEG_BIN, '-y',           # Sobrescreve saída sem perguntar
            '-f', 'concat',             # Usa o demuxer de concatenação
            '-safe', '0',               # Permite caminhos absolutos ou não sanitizados no arquivo de lista
            '-i', lista_txt_path,       # O arquivo de lista como entrada
            '-c', 'copy',               # Copia os codecs de áudio sem reencodar (muito rápido)
            arquivo_final               # O arquivo de saída final
        ]
        
        # O subprocess.run não mostra progresso para -c copy, então não passamos total_duration
        process = subprocess.run(comando, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        
        if process.returncode != 0:
            print(f"\n❌ Erro durante a unificação de áudio (código {process.returncode}):")
            print(process.stderr.decode(errors='ignore'))
            return False

        print(f"✅ Unificação de áudio concluída com sucesso para: {os.path.basename(arquivo_final)}")
        return True

    except FileNotFoundError:
        print(f"❌ Comando '{FFMPEG_BIN}' não encontrado. Certifique-se de que o FFmpeg está instalado no ambiente.")
        return False
    except subprocess.CalledProcessError as e:
        print(f"❌ Erro ao unificar áudio com FFmpeg: {e.stderr.decode(errors='ignore')}")
        return False
    except Exception as e:
        print(f"❌ Erro inesperado durante a unificação de áudio: {str(e)}")
        print(traceback.format_exc())
        return False
    finally:
        # Remove o arquivo de lista temporário
        if os.path.exists(lista_txt_path):
            try:
                os.remove(lista_txt_path)
                print(f"🧹 Arquivo de lista temporário removido: {os.path.basename(lista_txt_path)}")
            except Exception as e_unlink:
                print(f"⚠️ Não foi possível remover o arquivo de lista temporário {os.path.basename(lista_txt_path)}: {e_unlink}")


async def perform_conversion_task(file_path: str, voice: str, task_id: str, use_gemini_enhancement: bool = False, book_title: str = None):
    try:
        conversion_tasks[task_id].update({"status": "extracting", "message": "Iniciando extração de texto...", "progress": 0})
        text = await get_text_from_file(file_path, task_id)

        if not text:
            conversion_tasks[task_id].update({"status": "failed", "message": "Não foi possível extrair texto do arquivo."})
            return
        
        conversion_tasks[task_id].update({"status": "formatting", "message": "Formatando texto para melhor leitura TTS (Python nativo)...", "progress": 55})
        text_formatted = formatar_texto_para_tts(text)

        if use_gemini_enhancement and GEMINI_API_KEY:
            conversion_tasks[task_id].update({"status": "ai_enhancing", "message": "Revisando e melhorando texto com IA Gemini...", "progress": 57})
            print(f"Iniciando melhoria de texto com IA Gemini para tarefa {task_id}...")
            try:
                gemini_enhanced_text = await enhance_text_with_gemini(text_formatted)
                if gemini_enhanced_text.strip():
                    text_formatted = gemini_enhanced_text
                    print(f"Texto melhorado com Gemini para tarefa {task_id}.")
                else:
                    print(f"Gemini retornou texto vazio ou inválido para tarefa {task_id}. Usando texto original formatado.")
            except Exception as e_gemini:
                print(f"Erro ao usar Gemini para tarefa {task_id}: {e_gemini}. Prosseguindo com a formatação Python.")
        elif use_gemini_enhancement and not GEMINI_API_KEY:
            print(f"Usuário solicitou Gemini, mas a Chave API não está configurada para tarefa {task_id}. Prosseguindo com a formatação Python.")

        if not text_formatted.strip():
            conversion_tasks[task_id].update({"status": "failed", "message": "Texto vazio após formatação (ou IA). Nenhuma leitura TTS possível."})
            return

        # === Lógica para nomear o arquivo MP3 com o título do livro ===
        if book_title and book_title.strip():
            base_filename_clean = _limpar_nome_arquivo(book_title)
            original_filename_stem = _limpar_nome_arquivo(os.path.splitext(os.path.basename(file_path))[0])
            if not base_filename_clean:
                 final_audio_name_base = original_filename_stem
            else:
                 final_audio_name_base = f"{base_filename_clean}_{original_filename_stem[:20]}"
        else:
            final_audio_name_base = _limpar_nome_arquivo(os.path.splitext(os.path.basename(file_path))[0])
        
        audio_filename = f"{final_audio_name_base}.mp3"
        audio_filepath = os.path.join("audiobooks", audio_filename)
        conversion_tasks[task_id]["file_path"] = audio_filepath
        conversion_tasks[task_id]["total_characters"] = len(text_formatted)

        print(f"Iniciando geração de áudio com Edge TTS (Voz: {voice}) para {len(text_formatted)} caracteres formatados...")
        conversion_tasks[task_id].update({"status": "converting", "message": "Convertendo texto em áudio...", "progress": 60})

        LIMITE_CARACTERES_CHUNK_TTS = 5000
        CONCURRENCY_LIMIT = 2

        text_chunks = []
        current_chunk = ""
        paragraphs = text_formatted.split('\n\n')
        for p in paragraphs:
            if not p.strip():
                continue
            if len(current_chunk) + len(p) + 2 <= LIMITE_CARACTERES_CHUNK_TTS:
                current_chunk += (("\n\n" if current_chunk else "") + p)
            else:
                if current_chunk:
                    text_chunks.append(current_chunk)
                current_chunk = p
        if current_chunk:
            text_chunks.append(current_chunk)

        if not text_chunks:
            conversion_tasks[task_id].update({"status": "failed", "message": "Nenhuma parte de texto válida para conversão após divisão."})
            return

        total_chunks = len(text_chunks)
        temp_audio_chunk_paths = []
        
        semaphore = asyncio.Semaphore(CONCURRENCY_LIMIT)
        all_tts_tasks = []

        temp_chunks_dir = os.path.join("audiobooks", f"chunks_{task_id}")
        os.makedirs(temp_chunks_dir, exist_ok=True)


        async def convert_chunk_and_save_with_retry(chunk_text, voice_id, chunk_index, max_retries=3):
            chunk_temp_file = os.path.join(temp_chunks_dir, f"chunk_{chunk_index:04d}.mp3")
            
            for attempt in range(max_retries):
                try:
                    async with semaphore:
                        progress_tts = int(60 + (chunk_index / total_chunks) * 35)
                        conversion_tasks[task_id].update({"progress": progress_tts, "message": f"Gerando áudio (Parte {chunk_index+1}/{total_chunks}, Tentativa {attempt+1})..."})

                        communicate = edge_tts.Communicate(chunk_text, voice_id)
                        
                        await communicate.save(chunk_temp_file)
                        
                        if os.path.exists(chunk_temp_file) and os.path.getsize(chunk_temp_file) > 100:
                            print(f"✅ Chunk {chunk_index+1}/{total_chunks} concluído com sucesso e salvo em: {os.path.basename(chunk_temp_file)}")
                            return chunk_temp_file
                        else:
                            print(f"⚠️ Chunk {chunk_index+1}/{total_chunks}: Arquivo de áudio temporário inválido/vazio (tamanho: {os.path.getsize(chunk_temp_file) if os.path.exists(chunk_temp_file) else 0} bytes) (tentativa {attempt+1}).")
                            os.remove(chunk_temp_file) if os.path.exists(chunk_temp_file) else None
                            raise edge_tts.exceptions.NoAudioReceived

                except edge_tts.exceptions.NoAudioReceived:
                    print(f"❌ Chunk {chunk_index+1}/{total_chunks}: Sem áudio recebido. Retentando...")
                except asyncio.TimeoutError:
                    print(f"❌ Chunk {chunk_index+1}/{total_chunks}: Timeout na comunicação TTS. Retentando...")
                except Exception as e_chunk:
                    print(f"❌ Erro inesperado no chunk {chunk_index+1}/{total_chunks} (tentativa {attempt+1}): {type(e_chunk).__name__} - {e_chunk}")
                    print(traceback.format_exc())
                    os.remove(chunk_temp_file) if os.path.exists(chunk_temp_file) else None

                if attempt < max_retries - 1:
                    await asyncio.sleep(2 ** attempt)
            
            print(f"❌ Falha definitiva no chunk {chunk_index+1}/{total_chunks} após {max_retries} tentativas.")
            return None

        for i, chunk_text in enumerate(text_chunks):
            task = asyncio.create_task(convert_chunk_and_save_with_retry(chunk_text, voice, i))
            all_tts_tasks.append(task)

        results = await asyncio.gather(*all_tts_tasks, return_exceptions=True)

        successful_chunk_files = [res for res in results if isinstance(res, str) and os.path.exists(res) and os.path.getsize(res) > 100]
        
        if not successful_chunk_files:
            conversion_tasks[task_id].update({"status": "failed", "message": "Nenhum áudio válido foi gerado para o audiobook."})
            return

        conversion_tasks[task_id].update({"status": "merging_audio", "message": "Unificando partes do áudio...", "progress": 98})
        print(f"Unificando {len(successful_chunk_files)} arquivos de áudio temporários com FFmpeg...")
        
        if _unificar_audios_ffmpeg(successful_chunk_files, audio_filepath):
            print(f"Áudio final para tarefa {task_id} gerado e salvo em {audio_filepath}.")
            conversion_tasks[task_id].update({"status": "completed", "message": "Audiobook pronto para download!", "progress": 100})
        else:
            conversion_tasks[task_id].update({"status": "failed", "message": "Falha ao unificar partes do áudio. O audiobook pode estar incompleto."})
            print(f"❌ Falha ao unificar os áudios para tarefa {task_id}.")
            return

    except Exception as e:
        print(f"Erro na conversão da tarefa {task_id}: {e}")
        print(traceback.format_exc())
        conversion_tasks[task_id].update({"status": "failed", "message": f"Erro na conversão: {str(e)}"})
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Arquivo de texto original temporário {os.path.basename(file_path)} removido.")
        
        if os.path.exists(temp_chunks_dir):
            for temp_chunk_file in os.listdir(temp_chunks_dir):
                try:
                    os.remove(os.path.join(temp_chunks_dir, temp_chunk_file))
                except Exception as e_clean_chunk:
                    print(f"⚠️ Erro ao remover chunk temporário '{temp_chunk_file}': {e_clean_chunk}")
            os.rmdir(temp_chunks_dir)
            print(f"🧹 Diretório de chunks temporários removido: {os.path.basename(temp_chunks_dir)}")